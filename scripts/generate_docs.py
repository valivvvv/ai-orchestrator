"""
Generează documente DOCX fake pentru RAG demo.
"""
import random
import sys
from datetime import datetime, timedelta
from pathlib import Path

DOCS_DIR = Path(__file__).parent.parent / "data" / "documents"

COMPANIES = [
    ("TechSoft SRL", "ion.popescu@techsoft.ro", "0721-123-456", "IT & Software"),
    ("DataPro SA", "maria.ionescu@datapro.ro", "0722-234-567", "Consultanță DB"),
    ("CloudNet SRL", "andrei.stan@cloudnet.ro", "0723-345-678", "Cloud & Hosting"),
    ("WebDev Plus", "elena.marin@webdev.ro", "0724-456-789", "Web Development"),
    ("SecureIT SRL", "mihai.tudor@secureit.ro", "0725-567-890", "Cybersecurity"),
]

SERVICES = [
    ("Dezvoltare modul autentificare", 5200, 8500),
    ("Mentenanță lunară aplicație", 2500, 4500),
    ("Dezvoltare API REST", 7000, 12000),
    ("Consultanță PostgreSQL", 6000, 10000),
    ("Audit securitate infrastructură", 8000, 15000),
]


def generate_client_profile(company: str, email: str, phone: str, domain: str) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading(f'Profil Client: {company}', 0)

    doc.add_heading('Informații Contact', level=1)
    doc.add_paragraph(f'Denumire oficială: {company}')
    doc.add_paragraph(f'Email principal: {email}')
    doc.add_paragraph(f'Telefon: {phone}')
    doc.add_paragraph(f'Domeniu activitate: {domain}')

    doc.add_heading('Istoric Colaborare', level=1)
    projects = random.randint(3, 12)
    total_value = random.randint(30000, 200000)
    doc.add_paragraph(f'Număr proiecte finalizate: {projects}')
    doc.add_paragraph(f'Valoare totală colaborare: {total_value:,} RON')

    filename = f"client_{company.lower().replace(' ', '_').replace('.', '')}.docx"
    path = DOCS_DIR / filename
    doc.save(path)
    return path


def generate_invoice(invoice_id: int, company: str) -> Path:
    from docx import Document

    doc = Document()
    service, min_amt, max_amt = random.choice(SERVICES)
    amount = random.randint(min_amt, max_amt)
    date = datetime(2024, random.randint(1, 6), random.randint(1, 28))

    doc.add_heading('FACTURĂ FISCALĂ', 0)
    doc.add_paragraph(f'Serie: SL  Număr: {invoice_id:04d}')
    doc.add_paragraph(f'Data emiterii: {date.strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Beneficiar: {company}')

    doc.add_heading('Servicii', level=1)
    doc.add_paragraph(f'{service}: {amount:,} RON')

    tva = int(amount * 0.19)
    doc.add_paragraph(f'TVA 19%: {tva:,} RON')
    doc.add_paragraph(f'TOTAL: {amount + tva:,} RON')

    company_short = company.lower().split()[0]
    filename = f"factura_{invoice_id:04d}_{company_short}.docx"
    path = DOCS_DIR / filename
    doc.save(path)
    return path


def main():
    try:
        from docx import Document
    except ImportError:
        print("pip install python-docx")
        sys.exit(1)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    files = []

    # 5 profile
    print("Generez profile clienți...")
    for company, email, phone, domain in COMPANIES:
        f = generate_client_profile(company, email, phone, domain)
        files.append(f)
        print(f"  ✓ {f.name}")

    # 10 facturi
    print("Generez facturi...")
    for i in range(10):
        company = random.choice(COMPANIES)[0]
        f = generate_invoice(i + 1, company)
        files.append(f)
        print(f"  ✓ {f.name}")

    print(f"\n✓ {len(files)} documente în {DOCS_DIR}")


if __name__ == "__main__":
    main()
